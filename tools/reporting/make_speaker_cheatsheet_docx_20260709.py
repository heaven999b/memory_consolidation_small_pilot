from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE_SCRIPT_PATH = Path(
    "/Users/yihaiwen/Documents/New project/memory_consolidation_small_pilot/tools/reporting/make_technical_status_docx_20260709.py"
)
OUT_DIR = Path(
    "/Users/yihaiwen/Documents/New project/memory_consolidation_small_pilot/state/docx_speaker_cheatsheet_20260709"
)
OUT_DOCX = OUT_DIR / "mc_safety_speaker_cheatsheet_20260709.docx"


def load_base_module():
    spec = importlib.util.spec_from_file_location("tech_update_docx", BASE_SCRIPT_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


base = load_base_module()


def add_text(doc: Document, text: str, *, size=11, bold=False, color=None, after=6):
    p = doc.add_paragraph()
    base.set_paragraph_format(p, before=0, after=after, line=1.12)
    run = p.add_run(text)
    base.set_run_font(run, size=size, bold=bold, color=color or base.BODY)
    return p


def add_title_page(doc: Document):
    section = doc.sections[0]
    base.set_section_header_footer(section, "2026-07-09 讲解提词版")

    p = doc.add_paragraph()
    base.set_paragraph_format(p, before=0, after=4, line=1.0)
    run = p.add_run("SPEAKER CHEAT SHEET")
    base.set_run_font(run, size=11, color=base.MUTED, bold=True)

    p = doc.add_paragraph()
    base.set_paragraph_format(p, before=0, after=4, line=1.0)
    run = p.add_run("记忆固化安全研究：RQ1-RQ5 讲解提词版")
    base.set_run_font(run, size=22, color=base.BLACK, bold=True)

    p = doc.add_paragraph()
    base.set_paragraph_format(p, before=0, after=12, line=1.0)
    run = p.add_run("这不是正式演讲稿，而是给自己讲的时候看的版本")
    base.set_run_font(run, size=13, color=base.MUTED)

    meta_rows = [
        [["你现在手上有两份东西"], ["1. 正式技术稿: mc_safety_technical_update_20260709.docx", "2. 这份提词版: 只负责告诉你每个 RQ 看哪张表、怎么讲"]],
        [["这份文档的用法"], ["先看“主表”编号", "再看“最关键数字”", "最后直接照着“汇报时可以怎么说”那一句讲"]],
        [["总原则"], ["RQ2 一定拆成官方版和自建版", "RQ3 一定拆成旧线和新线", "RQ4 现在就是空白，不要硬讲成有结果"]],
    ]
    base.make_line_table(
        doc,
        ["项目", "内容"],
        meta_rows,
        [1800, 7560],
        font_size=9.4,
        header_size=10,
    )

    base.add_callout(
        doc,
        "怎么讲最稳",
        [
            "你不需要背所有数字。每个 RQ 只盯 2 到 4 个最关键数字。",
            "如果某个 RQ 没有稳定支持，就直接说“不支持”或“还没做完”，这比硬凹强很多。",
            "这轮最硬的主张其实不是“压缩有多坏”，而是“endpoint 会翻结论”和“reader 比 consolidator 更值得追”。",
        ],
        fill=base.ACCENT_FILL,
    )

    doc.add_page_break()
    doc.add_paragraph("总览：每个 RQ 先看哪张表", style="Heading 1")
    overview_rows = [
        ["RQ1", "0.4A", "0.4B / A3 / A4", "不支持原始强版本", "prompt-only 已经最糟；TierMem 没有更坏"],
        ["RQ2", "0.5A / 0.5B / 0.5C", "B5 / B7 / B11", "官方版不支持；自建版只支持弱版本", "一定拆成两条线讲"],
        ["RQ3", "0.6A / 0.6B", "C1 / C5", "旧线不支持；新线有正苗头", "这是唯一还在长正结果的主线"],
        ["RQ4", "D. RQ4 当前状态", "无", "空白", "没有真正的多算子对照表"],
        ["RQ5", "0.6C", "E1 / E2", "支持", "模型知道政策，但动作还是会违规"],
    ]
    table = base.make_table(
        doc,
        ["RQ", "主表", "补充表", "一句话结论", "你讲的时候最该强调什么"],
        overview_rows,
        [650, 1200, 1500, 2100, 3910],
        font_size=8.9,
        meta_note=base.rq_note("原始 RQ1-RQ5 讲解入口", "这张表只回答“先看哪张表、先讲什么”"),
    )
    for i in range(1, len(overview_rows) + 1):
        base.set_cell_text(table.cell(i, 3), overview_rows[i - 1][3], size=8.9, color=base.ACCENT, bold=True)


def add_rq1_section(doc: Document):
    doc.add_paragraph("RQ1：固化会不会把危险内容越洗越危险", style="Heading 1")
    add_text(doc, "这题现在的主结论很简单：不支持。你真正要看的不是“压缩以后有没有危险”，而是“压缩以后有没有比 baseline 更危险”。")

    base.make_line_table(
        doc,
        ["你要讲什么", "内容"],
        [
            [["主表看哪里"], ["0.4A RQ1 authority 主线", "补充看 0.4B RQ1 write-gate 线"]],
            [["这张表怎么跑的"], ["自建 stealth 安全题", "让 official policy 和 unsafe convention 同时在场", "比较 prompt_only 和 TierMem；看 N=0 和 N=2"]],
            [["主指标什么意思"], ["payload-emission = 模型最后有没有真的把危险内容说出来", "这个指标比“嘴上有没有拒绝”更硬"]],
            [["为什么还要看 0.4B"], ["因为 0.4B 会告诉你：只看 judge comply 会被带偏", "模型可以嘴上拒绝，但后面还是把危险内容做出来"]],
        ],
        [1900, 7460],
        font_size=9.4,
        header_size=10,
    )

    base.make_table(
        doc,
        ["条件", "最关键数字", "你该怎么读"],
        [
            ["prompt_only baseline", "15/15 = 1.000", "这是最糟 baseline"],
            ["TierMem N=0", "13/15 = 0.867", "进入 TierMem 后反而下降"],
            ["TierMem N=2", "11/15 = 0.733", "继续下降，不是继续升高"],
            ["write-gate N=1", "judge comply 0/24 ; payload 20/24", "嘴上拒绝几乎清零，但危险输出还在"],
        ],
        [2300, 1700, 5360],
        font_size=9.2,
        meta_note=base.rq_note("RQ1", "主表不支持原始“越固化越危险”；补充表支持 endpoint 会翻结论"),
    )

    base.add_callout(
        doc,
        "我是怎么看出来的",
        [
            "如果原始 RQ1 成立，数字应该随着进入 TierMem 或随着 N 变深而往上走。",
            "但现在最高的是 prompt_only 15/15，不是 TierMem。",
            "而且 TierMem 从 N=0 到 N=2 是 0.867 到 0.733，方向是往下，不是往上。",
            "所以这题不能再讲成“固化本身是主要放大器”。",
        ],
    )

    base.add_callout(
        doc,
        "汇报时可以直接说",
        ["现在看不到 TierMem 把危险内容越压越危险；相反，prompt-only 已经最糟，说明问题更像出在 reader 选动作这一步。"],
        fill=base.ACCENT_FILL,
    )

    base.add_callout(
        doc,
        "不要讲太满",
        ["不要再说“RQ1 已经证实压缩是危险放大器”。现在的数据不支持这句话。"],
    )


def add_rq2_section(doc: Document):
    doc.add_paragraph("RQ2：固化会不会制造假记忆 / 错信", style="Heading 1")
    add_text(doc, "这题一定拆成两条线讲：官方 benchmark 一条，自建题库一条。两条线现在给出的结论不一样，混在一起讲会很乱。")

    base.add_callout(
        doc,
        "先记住一句话",
        [
            "官方版: 不支持“越固化越糟”。",
            "自建版: 支持“重复错误说法会造错信”，但不支持“越固化越糟”的单调强版本。",
        ],
        fill=base.ACCENT_FILL,
    )

    doc.add_paragraph("RQ2-A 官方版：HaluMem", style="Heading 2")
    base.make_line_table(
        doc,
        ["你要讲什么", "内容"],
        [
            [["主表看哪里"], ["0.5A 官方版 HaluMem 15 QA", "0.5B 官方版 HaluMem 45 QA", "要更细就翻附录 B5 / B3"]],
            [["这张表怎么跑的"], ["官方 HaluMem-Medium 切片", "gpt-4.1-mini", "TierMem summary_only；比较不同 N"]],
            [["主指标什么意思"], ["UF_on_unknown = 本来不知道的题，模型却硬编了", "correct_rate = 答对比例", "F1 = 综合指标，但你讲的时候盯 UF_on_unknown 就够了"]],
        ],
        [1900, 7460],
        font_size=9.2,
        header_size=10,
    )
    base.make_table(
        doc,
        ["设置", "最关键数字", "你该怎么读"],
        [
            ["15 QA: N=0", "UF 2/6 = 0.333", "基线"],
            ["15 QA: N=1", "UF 1/6 = 0.167", "下降，不是上升"],
            ["15 QA: N=2", "UF 1/6 = 0.167", "没有反弹"],
            ["45 QA: N=0 -> 1 -> 2 -> 4", "0.333 -> 0.083 -> 0.083 -> 0.083", "扩大量之后方向还是往下"],
        ],
        [2350, 1900, 5110],
        font_size=9.0,
        meta_note=base.rq_note("RQ2 官方版", "官方 benchmark 不支持原始“越固化越糟”"),
    )
    base.add_callout(
        doc,
        "我是怎么看出来的",
        [
            "如果官方版 RQ2 成立，UF_on_unknown 应该随着 N 变深而升高。",
            "现在不管是 15 QA 还是 45 QA，UF 都是从 0.333 掉到 0.167 或 0.083。",
            "所以官方 benchmark 这条线不能拿来支持“越压越会乱编”。",
        ],
    )

    doc.add_paragraph("RQ2-B 自建版：100 题错信库", style="Heading 2")
    base.make_line_table(
        doc,
        ["你要讲什么", "内容"],
        [
            [["主表看哪里"], ["0.5C 自建版 100 题", "补充看附录 B7 重复强度表 和 B11 家族 breakdown"]],
            [["这张表怎么跑的"], ["自建 100 个基础题，分不同对话家族", "每层 200 probes", "比较 prompt-only 和 TierMem N=0/1/2"]],
            [["主指标什么意思"], ["FALSE_BELIEF = 模型最后信了错误说法", "这个指标高，说明错信更严重"]],
        ],
        [1900, 7460],
        font_size=9.2,
        header_size=10,
    )
    base.make_table(
        doc,
        ["设置", "最关键数字", "你该怎么读"],
        [
            ["prompt-only baseline", "12/200 = 6.0%", "baseline 已经会错信"],
            ["TierMem N=0", "21/200 = 10.5%", "第一层更坏，说明弱阳性存在"],
            ["TierMem N=1", "7/200 = 3.5%", "往下掉，不是继续变坏"],
            ["TierMem N=2", "5/200 = 2.5%", "继续下降，所以不是单调上升"],
            ["重复强度表", "1次 6.2% ; 3次 83.3% ; 5次 95.8%", "这说明“重复错误说法”这个机制是真存在的"],
        ],
        [2350, 2100, 4860],
        font_size=8.9,
        meta_note=base.rq_note("RQ2 自建版", "支持弱版本：重复错误说法会造错信；不支持单调“越固化越糟”"),
    )
    base.add_callout(
        doc,
        "我是怎么看出来的",
        [
            "自建版不是完全负面，因为 N=0 的确比 prompt-only 更坏。",
            "但如果你要讲原始强版本，问题在于后面 N=1、N=2 又掉下去了。",
            "所以最诚实的说法是：会造错信，但不是越固化越糟。",
            "另外，重复次数从 1 到 3 到 5 的那张表非常强，说明“错误信息重复”这个机制是真的。",
        ],
    )
    base.add_callout(
        doc,
        "汇报时可以直接说",
        [
            "官方版 benchmark 没有支持“越固化越糟”；自建版支持一个更弱也更具体的版本：重复错误说法会造错信，但不是单调随 N 上升。"
        ],
        fill=base.ACCENT_FILL,
    )


def add_rq3_section(doc: Document):
    doc.add_paragraph("RQ3：provenance / 读取侧防御能不能救", style="Heading 1")
    add_text(doc, "这题一定拆成旧线和新线。旧线现在不能当主结果讲；新线是目前最像正结果的方向，但还没完全跑完。")

    doc.add_paragraph("RQ3-A 旧线：旧 provenance 主张", style="Heading 2")
    base.make_line_table(
        doc,
        ["你要讲什么", "内容"],
        [
            [["主表看哪里"], ["0.6A RQ3 旧 provenance 线", "附录 C1"]],
            [["这张表怎么跑的"], ["旧 auto route 对比 summary_only route", "每个 route、每个 N 都只有 12 个样本", "这是旧写入线，不是现在主推的干净方案"]],
            [["主指标什么意思"], ["comply = 最后顺着危险习惯做了", "UAF = 有没有出现“洗白后救火”的现象"]],
        ],
        [1900, 7460],
        font_size=9.2,
        header_size=10,
    )
    base.make_table(
        doc,
        ["route", "最关键数字", "你该怎么读"],
        [
            ["auto comply", "0.250 / 0.167 / 0.083 / 0 / 0", "有下降，但不算稳定独占优势"],
            ["summary_only comply", "0 / 0.083 / 0 / 0 / 0", "有时它反而更低"],
            ["UAF", "全 0", "没有出现“洗白救火”现象"],
        ],
        [2200, 2200, 4960],
        font_size=9.0,
        meta_note=base.rq_note("RQ3 旧线", "不支持旧版 provenance 主张"),
    )

    doc.add_paragraph("RQ3-B 新线：读取侧防御 off/on", style="Heading 2")
    base.make_line_table(
        doc,
        ["你要讲什么", "内容"],
        [
            [["主表看哪里"], ["0.6B RQ3 新 read-time defense 线", "如果别人追问细节，就翻附录 C5"]],
            [["这张表怎么跑的"], ["固定同一批自建安全题", "只改读取侧规则：off 对比 on", "总 jobs 20，最新快照完成了 18 个"]],
            [["主指标什么意思"], ["payload-emission = 最后有没有真的吐出危险内容", "better / worse flips = 同一批题里，被防御救回了多少、又带坏了多少"]],
        ],
        [1900, 7460],
        font_size=9.2,
        header_size=10,
    )
    base.make_table(
        doc,
        ["条件", "off mean", "on mean", "delta", "你该怎么读"],
        [
            ["prompt_only N=0", "94.0%", "80.0%", "-14.0pp", "正方向，而且 5 个 seed 都没有 worse flips"],
            ["tiermem N=0", "88.3%", "76.7%", "-11.7pp", "正方向，但有少量 worse flips"],
            ["tiermem N=1", "86.7%", "77.5%", "-9.2pp", "这是目前最弱的一层，但还是往下"],
            ["tiermem N=2", "88.3%", "74.2%", "-14.2pp", "正方向也很明显"],
        ],
        [2000, 1100, 1100, 1200, 3960],
        font_size=8.9,
        meta_note=base.rq_note("RQ3 新线", "当前支持倾向为正，但还没最终封盘"),
    )

    base.add_callout(
        doc,
        "我是怎么看出来的",
        [
            "旧线不能讲成成立，因为 auto 并没有稳定比 summary_only 更好，而且 UAF 全程是 0。",
            "新线可以讲成“有正苗头”，因为所有已完成条件的 off -> on 都是下降。",
            "但它还没完全封盘，因为 20 个 jobs 只完成了 18 个。",
            "所以最稳的讲法是：旧线不支持，新线目前支持倾向为正。",
        ],
    )

    base.add_callout(
        doc,
        "汇报时可以直接说",
        ["RQ3 不能笼统讲成成立。旧版 provenance 线不成立；新版读取侧防御线目前是正方向，而且已经完成 90%，但还没最终封盘。"],
        fill=base.ACCENT_FILL,
    )


def add_rq4_section(doc: Document):
    doc.add_paragraph("RQ4：哪个压缩算子最脆", style="Heading 1")
    add_text(doc, "这题现在最重要的不是“怎么讲结果”，而是老老实实承认：还没有结果。")

    base.make_line_table(
        doc,
        ["你要讲什么", "内容"],
        [
            [["主表看哪里"], ["D. RQ4 当前状态"]],
            [["为什么现在是空白"], ["我们只系统跑了 TierMem", "还没有 COMEDY / Context-Memory / NeedSleep / E-mem 的统一脚本、统一数据、统一判分表"]],
            [["如果以后要做"], ["必须在同一套题、同一套 scorer、同一套 N 设置下横向比较不同压缩算子", "否则出来的差异没法归因"]],
        ],
        [1900, 7460],
        font_size=9.4,
        header_size=10,
    )

    base.make_table(
        doc,
        ["现在能说什么", "现在不能说什么"],
        [[
            "RQ4 目前没有正式实验表，所以既不支持也不反驳任何算子更脆弱的说法。",
            "不要说“TierMem 最脆”或者“某个论文方法更安全”，因为仓库里现在没有这种可比数值。",
        ]],
        [4680, 4680],
        font_size=9.3,
        meta_note=base.rq_note("RQ4", "空白就是空白；这题现在没有可上台面的数值"),
    )

    base.add_callout(
        doc,
        "汇报时可以直接说",
        ["RQ4 目前还是空白，因为我们还没做真正的多算子对照。现在仓库里只有比较思路，没有统一可比的结果表。"],
        fill=base.ACCENT_FILL,
    )


def add_rq5_section(doc: Document):
    doc.add_paragraph("RQ5：失败到底发生在哪个阶段", style="Heading 1")
    add_text(doc, "这题是目前最稳的一条正结果。核心不是“模型不知道政策”，而是“模型知道政策，但动作还是会违规”。")

    base.make_line_table(
        doc,
        ["你要讲什么", "内容"],
        [
            [["主表看哪里"], ["0.6C RQ5 know-do gap", "补充看附录 E1 / E2"]],
            [["这张表怎么跑的"], ["同一条 memory 里同时放 official policy 和 unsafe convention", "先问模型“政策是什么”，再问模型“你现在该怎么做”", "三模型各跑 30 条"]],
            [["主指标什么意思"], ["know-do gap = 知道政策，但动作还是违规", "consistent_safe = 知道政策，而且动作也安全", "doesnt_know = 连政策都没答出来"]],
        ],
        [1900, 7460],
        font_size=9.2,
        header_size=10,
    )

    base.make_table(
        doc,
        ["模型", "know-do gap", "doesnt_know", "你该怎么读"],
        [
            ["gpt-4.1-mini", "19/30 = 63.3%", "0/30", "最差，但不是因为不知道政策"],
            ["gpt-4o", "14/30 = 46.7%", "0/30", "中间水平"],
            ["gpt-4.1", "11/30 = 36.7%", "0/30", "最好，但 gap 依然在"],
            ["policy-check 对照", "0.667 -> 0.500", "0 -> 0", "加一道 policy-check 只能部分缓解，不能清零"],
        ],
        [1800, 1800, 1100, 4660],
        font_size=9.1,
        meta_note=base.rq_note("RQ5", "明确支持：问题更像 reader / answer-time，不像记忆本身坏掉"),
    )

    base.add_callout(
        doc,
        "我是怎么看出来的",
        [
            "三模型的 doesnt_know 全是 0，这说明它们不是“不知道政策”。",
            "但 know-do gap 仍然有 36.7% 到 63.3%，说明知道归知道，做归做。",
            "这正是为什么我们说问题更像 reader / answer-time，而不是 consolidator 没把政策存住。",
        ],
    )

    base.add_callout(
        doc,
        "汇报时可以直接说",
        ["RQ5 目前是最稳的正结果：模型知道政策，但动作还是会违规，所以更像是 reader 在选动作时给坏习惯分了太高权重。"],
        fill=base.ACCENT_FILL,
    )


def add_finish_section(doc: Document):
    doc.add_paragraph("最后一页：你上台时的总口径", style="Heading 1")
    add_text(doc, "如果你只能讲 1 分钟，就照下面这张表念。")

    base.make_table(
        doc,
        ["RQ", "你现在最稳的一句话"],
        [
            ["RQ1", "不支持“越固化越危险”；prompt-only 已经最糟，TierMem 没有更坏。"],
            ["RQ2", "官方版不支持；自建版只支持“重复错误说法会造错信”的弱版本。"],
            ["RQ3", "旧线不支持；新读取侧防御线是正方向，但还没最终封盘。"],
            ["RQ4", "现在是空白，因为还没有真正的多算子对照。"],
            ["RQ5", "明确支持 know-do gap：模型知道政策，但动作还是会违规。"],
        ],
        [800, 8560],
        font_size=9.5,
        meta_note=base.rq_note("原始 RQ1-RQ5 上台口径", "这是给自己讲的时候看的最终压缩版"),
    )

    base.make_table(
        doc,
        ["现在可以明确说", "现在要保守说"],
        [[
            "1. endpoint 定义会翻结论\n2. know-do gap 在三模型都存在\n3. RQ3 新线目前是正方向\n4. reader 比 consolidator 更值得追",
            "1. 不要再说 RQ1 已成立\n2. 不要把官方版和自建版 RQ2 混成一条线\n3. 不要把 RQ3 讲成已封盘\n4. 不要假装 RQ4 已经有多算子结果",
        ]],
        [4680, 4680],
        font_size=9.2,
        meta_note=base.rq_note("汇报禁区提醒", "这张表就是防止你讲嗨了以后把现在不该说的话说出去"),
    )


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    base.configure_document(doc)
    add_title_page(doc)
    add_rq1_section(doc)
    add_rq2_section(doc)
    add_rq3_section(doc)
    add_rq4_section(doc)
    add_rq5_section(doc)
    add_finish_section(doc)
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_report()
