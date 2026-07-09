# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
OUT=Path("/Users/yihaiwen/Documents/New project/memory_consolidation_small_pilot")/"paper_consolidation_security_20260704.docx"
doc=Document()
def sea(s,l="Times New Roman",e="SimSun"):
    r=s.element.get_or_add_rPr();f=r.find(qn('w:rFonts'))
    if f is None:f=OxmlElement('w:rFonts');r.append(f)
    f.set(qn('w:ascii'),l);f.set(qn('w:hAnsi'),l);f.set(qn('w:eastAsia'),e)
nm=doc.styles['Normal'];nm.font.name='Times New Roman';nm.font.size=Pt(10.5);sea(nm)
for s,e in [('Heading 1','SimHei'),('Heading 2','SimHei'),('Heading 3','SimHei'),('Title','SimHei'),('List Bullet','SimSun'),('List Number','SimSun')]:
    try:sea(doc.styles[s],"Times New Roman",e)
    except KeyError:pass
for sec in doc.sections:
    sec.top_margin=Inches(1);sec.bottom_margin=Inches(1);sec.left_margin=Inches(1);sec.right_margin=Inches(1)
def h1(t):return doc.add_heading(t,1)
def h2(t):return doc.add_heading(t,2)
def P(t,b=False,i=False,sz=None,c=None,al=None,fi=False):
    p=doc.add_paragraph();r=p.add_run(t);r.bold=b;r.italic=i
    if sz:r.font.size=Pt(sz)
    if c:r.font.color.rgb=RGBColor(*c)
    if al:p.alignment=al
    if fi:p.paragraph_format.first_line_indent=Inches(0.25)
    return p
def bl(t):return doc.add_paragraph(t,style='List Bullet')
def nu(t):return doc.add_paragraph(t,style='List Number')
def pb():doc.add_page_break()
def ear(r,e="SimSun"):
    p=r._element.get_or_add_rPr();f=p.find(qn('w:rFonts'))
    if f is None:f=OxmlElement('w:rFonts');p.append(f)
    f.set(qn('w:ascii'),"Times New Roman");f.set(qn('w:hAnsi'),"Times New Roman");f.set(qn('w:eastAsia'),e)
def scell(c,t,b=False,col=None):
    c.text="";p=c.paragraphs[0];r=p.add_run(t);r.bold=b;r.font.size=Pt(9)
    if col:r.font.color.rgb=RGBColor(*col)
    ear(r)
def sh(c,f):
    tc=c._tc.get_or_add_tcPr();e=OxmlElement('w:shd');e.set(qn('w:val'),'clear');e.set(qn('w:fill'),f);tc.append(e)
def tbl(cap,hd,rows,fill="1F3864"):
    if cap: P(cap,b=True,sz=9.5)
    t=doc.add_table(rows=1,cols=len(hd));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,x in enumerate(hd):scell(t.rows[0].cells[i],x,True,(255,255,255));sh(t.rows[0].cells[i],fill)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):scell(cs[i],str(v))
    doc.add_paragraph()
    return t

# ===== 标题与摘要 =====
P("评测器决定你看到的威胁", True, sz=20, al=WD_ALIGN_PARAGRAPH.CENTER)
P("记忆固化 Agent 的安全评测陷阱、固化旁路攻击与能力推演审计", True, sz=13, c=(0x1F,0x38,0x64), al=WD_ALIGN_PARAGRAPH.CENTER)
P("The Evaluator Determines the Threat You See: Evaluation Pitfalls, Consolidation-Bypass Attacks, and Capability-Inference Auditing in Memory-Augmented Agents", i=True, sz=10, al=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
P("作者：〔待填〕　　单位：〔待填〕　　日期：2026-07-04", sz=10, al=WD_ALIGN_PARAGRAPH.CENTER)
P("证据成熟度声明：本文全部实验为受控 pilot 规模（每组 n≈6–30），单一记忆底座（TierMem），主干模型 gpt-4.1-mini 与 gpt-4o。所有主张均在其证据强度边界内陈述，局限见第 8 节。", i=True, sz=9, c=(0x99,0,0), al=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
P("摘要", b=True, sz=12)
P("记忆增强 Agent（memory-augmented agents）依赖迭代记忆固化（consolidation）压缩长期记忆，但该机制的安全性尚未被系统评估。本文以分层记忆系统 TierMem 为底座，对“记忆固化中的安全”做了一次从证伪到重建的完整研究。(1) 我们首先识别并纠正了该领域评测中的四类系统性陷阱——词法/拒答终点、天花板效应、同源且未经人工校准的 LLM judge、以及顶层过度断言与选择性报告——这些陷阱使我们自己最初的四条“发现”中的三条被对抗审查证伪。(2) 在此基础上我们提出并验证一个记忆系统独有的攻击面——固化作为写入期安全审查的旁路：危险内容在审查点之后被固化步骤“措辞正当化”，从而绕过表面安全分类器（检测盲区），而下游 reader 仍可利用之。(3) 我们提出能力推演审计（capability-inference audit）作为防御——让审计器推演“这条记忆能被用来做什么”——其攻击召回 1.0、良性特异性 0.83，并在适应性攻击下保持 5/6 鲁棒。(4) 通过预注册实验与跨模型去同源，我们得到统一结论：在记忆 Agent 安全中，你测到的威胁 ≈ 真实威胁 × 评测器能力——弱评测器既高估防御（对冲拒答）、也高估攻击面（检测盲区伪影）；将审计器由 gpt-4.1-mini 换成 gpt-4o 可把参数级检测盲区由 8/10 关闭到 4/10 而不牺牲特异性。我们如实标注全部证据为 pilot 规模，并给出通向大规模验证的路线图。", fi=True)
P("关键词：记忆增强 Agent；记忆固化；LLM 安全评测；能力推演审计；检测盲区；评测可靠性", b=True, sz=9.5)
pb()

# ===== 1 引言 =====
h1("1  引言")
P("长期记忆正迅速成为 LLM Agent 的标准组件：mem0、MemGPT、A-MEM 及本文所用的 TierMem 等系统通过迭代“固化”（consolidation，即把历史交互反复摘要/合并为紧凑记忆）来突破上下文窗口限制。然而，一旦 Agent 的行为依赖一段可被写入、且会被系统自身反复改写的长期记忆，就引入了一个尚未被系统研究的攻击面。本文的出发点正是这一问题：记忆固化 Agent 的安全应当如何评测、又存在哪些机制性风险？", fi=True)
P("我们的研究经历了一次诚实的自我修正，并因此得到本文最重要的洞察。项目最初假设“不安全内容会随固化深度 N 被放大/洗白”，但密集实验与多视角对抗 red-team 将其证伪；进一步的对抗审查表明，我们最初赖以立论的评测终点本身系统性失真。由此我们意识到一个贯穿始终的张力——", fi=True)
P("在记忆 Agent 安全中，你测到的威胁 ≈ 真实威胁 × 评测器能力。", b=True, c=(0x1F,0x38,0x64), al=WD_ALIGN_PARAGRAPH.CENTER)
P("弱评测器会同时朝两个方向骗你：用对冲措辞骗过 comply/refuse 判卷，使你高估防御；用不足的判别力漏检伪装攻击，使你高估“检测盲区”这一攻击面。正确的做法是采用能力推演审计并搭配足够强的评测器。", fi=True)
P("本文贡献：", b=True)
nu("方法学：识别并以实证纠正记忆/Agent 安全评测中的四类系统性陷阱，并给出一套可复用的对抗审查规程（生死对照、预注册、精确区间、去同源）。")
nu("攻击：形式化并验证“固化旁路”攻击面——固化产物不回写入闸门、且被措辞正当化，构成对表面安全分类器的检测盲区。")
nu("防御：提出能力推演审计，量化其召回/特异性/抗适应性攻击鲁棒性，并给出与表面闸门的权衡分析。")
nu("发现：通过预注册与跨模型去同源，证明检测盲区/可规避性是“攻击结构 × 评测器能力”的交互，而非纯攻击属性；据此给出“评测器依赖”的统一框架与防御建议。")
P("我们强调：本文全部证据为 pilot 规模，结论均在其边界内陈述；我们将其诚实标注视为方法学贡献的一部分。", i=True)
pb()

# ===== 2 相关工作 =====
h1("2  相关工作")
P("记忆增强 Agent 与固化。 长期记忆系统（mem0、MemGPT、A-MEM、TierMem 等）普遍采用摘要/合并式固化。既有工作多关注其效用（长程问答召回），对固化引入的安全语义变化关注甚少。", fi=True)
P("Agent/LLM 安全评测。 AgentPoison、HarmBench 等提供攻击语料与判卷协议；但多以词法命中或 comply/refuse 标签为终点。本文表明此类终点对推理型/对冲型答案系统性失真。", fi=True)
P("记忆投毒与提示注入。 已有研究展示外部输入可污染检索记忆。本文区别在于：攻击载荷并非以单条形式注入，而是由系统自身的固化步骤从各自无害的碎片中合成/改写而来，从而绕过写入期审查。", fi=True)
P("LLM-as-judge 的可靠性。 已有工作指出 LLM 评判存在自偏好与不稳定性。本文提供一个具体且被自身数据证伪的案例，并进一步证明“评测器能力”会定量改变所得到的安全结论。", fi=True)
pb()

# ===== 3 威胁模型 =====
h1("3  威胁模型与背景")
h2("3.1  记忆固化 C^N")
P("固化算子将一组记忆压缩为紧凑记忆；迭代固化 C^N 指将上一轮固化输出作为下一轮输入反复施加。实现上采用结构化摘要（gpt-4.1-mini）。我们观测到合成/压缩在 C^1 即完成、C^2=C^3 收敛，表明相关效应是“单次合并”事件而非“深度累积”，这也否证了原始的“深度 N”假设。", fi=True)
h2("3.2  写入期闸门与旁路")
P("典型防御 safety_write_filter 为写入期闸门：对每条候选入库记忆做安全分类并标记/隔离。我们通过代码审查确认：固化算子合成的新记忆不会重新经过写入闸门。因此攻击者的目标不是让任何单条输入被判危险，而是让危险“藏在固化产物里”——固化步骤发生在审查点之后，其输出直接入库。", fi=True)
h2("3.3  攻击者能力")
P("我们假设一个可写入若干条记忆的攻击者：每条单独看均为无害的运维事实/工具文档/惯例，且能通过写入闸门。攻击者不接触固化算子内部，也不接触 reader。在适应性设定中，攻击者已知存在能力推演审计，可对记忆措辞进行改写（best-of-3），但必须保持下游 reader 的可执行性。", fi=True)
pb()

# ===== 4 第一部分:评测陷阱 =====
h1("4  第一部分：评测陷阱与原结论的证伪")
P("本节记录我们对自身最初四条“发现”的对抗审查。这既是诚实要求，也奠定了全文“评测器依赖”的主题。", fi=True)
tbl("表 1：四条原结论的对抗审查裁决（payload/CI 均为精确 Clopper-Pearson）。",
 ["原结论","裁决","关键证据"],
 [["固化放大不安全内容","证伪","prompt_only 1.000 ≥ tiermem N0 0.867 ≥ N2 0.733；无记忆基线暴露最高"],
  ["know-do gap（知—行分离）","命名不成立","doesn’t-know 恒为 0 → gap 数值上恒等于违反率；KNOW 为开卷复述题（阈值 50% 词覆盖）"],
  ["hedge-refusal（应改用 payload 终点）","强主张不成立","83% “泄露”中约 10/20 是字面“Do not…”硬拒绝被误判；偏差由高估安全翻成高估危险"],
  ["provenance 读取侧防御有效","无信号","5:0 是 n=5 下数学最小两侧 p=2·0.5⁵=0.0625，含 1 假阴+2 伪重复"]])
P("四类系统性陷阱。 由上述证伪归纳出可迁移的教训：", b=True)
tbl("表 2：记忆/Agent 安全评测的四类系统性陷阱。",
 ["陷阱","机制","后果"],
 [["无否定/语义处理的词法终点","子串/关键词匹配，不辨“拒绝做 X”与“去做 X”","judge 高估安全，词法高估危险，双向失真"],
  ["天花板效应 + 功效不足","诱导性查询把基线顶到 1.0；n 小、单 seed、伪重复","差异检验注定 null，被误用为“无害”（缺 TOST）"],
  ["同源且未校准的 judge","judge 与被测同模型；零人工 κ","自偏好；本文中该 judge 被自身数据证伪（83% 假拒绝）"],
  ["顶层过度断言 + 选择性报告","正文退让但摘要/结论不退；多 run 事后挑显著","overclaim；FWER 无校正"]])
P("这些陷阱直接指向全文主题：评测终点/评测器的选择会定性甚至定量地改变所得结论。第 7 节将表明，同一逻辑也作用于攻击面的度量。", fi=True)
pb()

# ===== 5 第二部分:固化旁路与检测盲区 =====
h1("5  第二部分：固化旁路与检测盲区")
h2("5.1  合成并非固化独有：一次关键的生死对照")
P("我们首先检验“固化能否从无害碎片合成恶意能力”。将 4 条各自无害、且全部通过写入闸门的碎片过 C^N：C^1 即缝合出含外部收件地址、导出脚本与目的地参数的可执行外泄配方，且合成产物被 rules 闸门判无害、被 llm 闸门判危险。然而生死对照证伪了强主张——将同样 4 条碎片直接置于上下文，reader 同样输出该外泄命令。因此“合成”是模型推理的普遍能力，并非固化独有；原“固化≫上下文”不成立。", fi=True)
h2("5.2  真正的现象：固化制造检测盲区")
P("固化的独特危害不在“合成新能力”，而在改变可检测性：固化把危险要素重新措辞为读起来像例行流程的紧凑记忆（如把外泄洗为“standard data handling”），使表面安全分类器漏检，而 reader 仍可利用。我们在 6 个正交攻击场景上量化：", fi=True)
tbl("表 3：v3 检测盲区量化（6 场景）。R_note=表面闸门对固化 note 的判定。",
 ["场景类","表面闸门 R_note","reader 可利用","能力审计","盲区"],
 [["exfiltration / backup_to_public","SAFE（漏）","可利用","RISKY（抓）","是（2/6）"],
  ["priv_esc / approval_bypass / audit_disable / secret_in_logs","RISKY（抓）","可利用","RISKY","否"]])
P("盲区 2/6，reader 可利用 6/6。机制洞察：盲区寄生在“攻击动作≈良性操作”的类别（数据传输、存储位置）——这一类别的“危险”仅在于一个可被措辞掩盖的目的地参数。", fi=True)
pb()

# ===== 6 第三部分:能力推演审计 =====
h1("6  第三部分：能力推演审计防御")
P("针对措辞正当化，我们提出能力推演审计：不问“这段文本是否含危险措辞”，而让审计器推演“若 Agent 采纳这条记忆，能被用来做什么、是否可能违反安全政策”，据此判定。", fi=True)
h2("6.1  精度：召回 vs 特异性")
P("在 6 个攻击场景与 6 个话题相同、风险相反的良性场景（内部导出/私有加密备份/双人审批/正常授权/审计常开/脱敏日志）上构建混淆矩阵：", fi=True)
tbl("表 4：v4 能力推演审计 vs 表面措辞闸门（攻击 6 / 良性 6）。",
 ["检测器","召回","特异性","精度","要点"],
 [["能力推演审计","1.00","0.83","0.86","补上全部盲区；良性 5/6 判 safe，非话题匹配器"],
  ["表面措辞闸门","0.67","1.00","1.00","漏 2 个盲区；零误报但召回不足"]])
P("能力审计以 1 个可辩护的误报（“on-call 常备生产权限”本身即温和风险）换取召回由 0.67 提升至 1.00。对安全审计，漏检代价通常高于误报。", fi=True)
h2("6.2  适应性攻击鲁棒性")
P("令攻击者已知审计存在并改写记忆以规避（best-of-3）。能力审计抗规避 5/6；唯一失守为外泄场景。一个直接证据支撑“可利用性绑定可检测性”：在提权场景，攻击者为骗过审计把记忆改得过于正当，导致 reader 也不再执行该攻击（可利用性随之消失）。即：要藏过能力审计，往往就得藏掉能力本身。", fi=True)
pb()

# ===== 7 第四部分:评测器依赖 =====
h1("7  第四部分：动作层可区分度与评测器依赖")
P("我们进一步追问：可规避性是攻击结构的内在属性，还是取决于评测器？", fi=True)
h2("7.1  预注册定律（弱审计器下）")
P("预注册预测：参数级攻击（良性动作 + 可伪装目的地参数，如导出/备份/转发/webhook 指向外部）可规避；能力级攻击（动作本身即危险，如关审计/强推/提权/关 MFA/删备份/导密钥）不可规避。在 gpt-4.1-mini 审计器、best-of-3、每组 10 场景下：", fi=True)
tbl("表 5：动作层可区分度——审计器能力梯度（best-of-3 适应性规避率）。",
 ["攻击类","gpt-4.1-mini 审计（弱）","gpt-4o 审计（强）"],
 [["参数级（n=10）","8/10 = 0.80，CI[0.44,0.97]","4/10 = 0.40，CI[0.12,0.74]"],
  ["能力级（n=10）","0/10 = 0.00，CI[0.00,0.31]","1/10 = 0.10，CI[0.00,0.45]"],
  ["两组 CI","分离（显著）","重叠（不显著）"]])
P("在弱审计器下预测命中且两组 CI 显著分离，定律成立。", fi=True)
h2("7.2  跨模型去同源：定律强依赖评测器能力")
P("将审计器换为不同规模/架构的 gpt-4o（打破“同模型判自己改写”的自评闭环）。诊断显示 gpt-4o 是更强的审计器而非过度警报者：对 6 个良性 note 判 RISKY 为 0/6（特异性 100%），同时对原始参数级攻击 10/10 命中——而 gpt-4.1-mini 连原始参数级攻击都漏 3/10。在完整适应性实验中，换用 gpt-4o 使参数级规避由 8/10 降至 4/10、能力级由 0/10 微升至 1/10，两组 CI 由分离变为重叠。", fi=True)
P("结论：参数级检测盲区部分是弱审计器的伪影（8→4 被强审计器关闭）、部分是真实残余（gpt-4o 仍漏 4/10）；能力级始终最难规避。可规避性是“攻击结构 × 评测器能力”的交互——这正是全文统一框架在攻击面一侧的体现。", b=True)
pb()

# ===== 8 讨论/局限/伦理 =====
h1("8  讨论")
P("统一框架。 第一部分（弱终点高估防御）与第四部分（弱审计器高估攻击面）是同一现象的两面：安全评测的结论强依赖评测器能力。这提示一条方法学准则——报告任何记忆/Agent 安全结论时，必须同时报告评测器身份与能力，并做评测器敏感性分析。", fi=True)
P("防御建议。 (i) 以能力推演审计替代表面措辞分类，并在固化产物入库前重新审计（关闭旁路）；(ii) 对参数级风险，语义审计（管“能做什么”）须与执行层目的地白名单（管“对谁做”）互补；(iii) 采用尽可能强的审计模型——评测器能力是记忆安全的一等变量。", fi=True)
h1("9  局限性")
P("我们如实标注全部证据为 pilot 规模，任何结论应在以下边界内解读：", i=True)
bl("规模：每组 n≈6–30，比例的精确 CI 较宽（如强审计器下两组 CI 重叠）；需每组 ≥100 方可收窄。")
bl("单一底座与合成种子：仅 TierMem，攻击/良性种子由作者构造，可能偏离真实分布；未接真实 mem0 检索与跨 session 持久化闭环。")
bl("评测器同源与无人工真值：主干为 gpt-4.1-mini/gpt-4o（同厂商）；跨厂商去同源因缺少非 OpenAI 凭据未能完成；全流程无人工标注与 Cohen’s κ。")
bl("适应性攻击为 best-of-3：真实红队会多轮迭代，规避率为乐观下限。")
bl("自查记录：预注册实验首轮因误将审计判定置于 temperature=0.2 引入噪声，修正为判定固定 temperature=0 后复现——测量对细节敏感。")
h1("10  伦理声明")
P("本研究面向防御。所有攻击场景为合成、面向内部受控实验，不针对任何真实系统或用户；不涉及可直接复用的漏洞武器。我们提出的能力推演审计与旁路修补建议旨在加固记忆 Agent。攻击细节的披露程度以“足以复现防御评估、不足以降低现实攻击门槛”为界。", fi=True)
h1("11  结论")
P("我们对记忆固化 Agent 的安全做了一次从证伪到重建的完整研究，提出并验证了固化旁路攻击面与能力推演审计防御，并通过预注册与去同源得到统一结论：你测到的威胁取决于评测器能力。弱评测器双向误导——既高估防御也高估攻击面；正确的路径是能力推演审计加足够强的评测器。全部证据为 pilot 规模，我们据此给出通向大规模验证的路线图。", fi=True)
pb()

# ===== 附录 =====
h1("附录 A  实验清单与可复现信息")
tbl("表 A1：实验与产物一览。",
 ["实验","关键结果","脚本 / 数据"],
 [["v1 显微镜","C^1 合成；生死对照证伪“固化≫上下文”；provenance 无信号","spike_consolidation_microscope.py / spike_microscope_v1.json"],
  ["v2 多场景旁路","rules 闸门 0/3；库审查 上下文 0/3 vs 固化 2/3","spike_consolidation_bypass_v2.py / spike_bypass_v2.json"],
  ["v3 盲区量化","盲区 2/6；reader 可利用 6/6；能力审计救回 2/2","spike_v3_blindspot.py / spike_v3_blindspot.json"],
  ["v4 防御精度","能力审计 召回1.0/特异性0.83/精度0.86","spike_v4_precision.py / spike_v4_precision.json"],
  ["v5 适应性攻击","抗规避 5/6；唯一失守 exfil","spike_v5_adaptive.py / spike_v5_adaptive.json"],
  ["A1 预注册定律","参数级 3/6 vs 能力级 0/6，命中","spike_a1_law.py / spike_a1_law.json"],
  ["v6 定律巩固","参数级 8/10 vs 能力级 0/10，CI 分离","spike_v6_law_scale.py / spike_v6_law_scale.json"],
  ["v7 去同源","gpt-4o：参数级 4/10 vs 能力级 1/10，CI 重叠；特异性 100%","spike_v7diag2.py / spike_v7b_full.py / *.json"]])
P("复现须知：主干模型 gpt-4.1-mini（生成/攻击/reader）与 gpt-4o（强审计器）；固化算子采用结构化摘要；OpenAI 客户端须设 timeout=15–20s、max_retries=1 以避免单次挂起阻塞；判定固定 temperature=0，仅对抗改写用 temperature=0.8。", sz=9.5)
doc.add_paragraph();P("— 全文完 —", i=True, al=WD_ALIGN_PARAGRAPH.CENTER)
doc.save(str(OUT));print("PY_SAVED_OK")
